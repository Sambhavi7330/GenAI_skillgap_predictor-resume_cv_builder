# ============================
# Gen-AI Skill Gap Predictor + Resume & CV Builder (Groq)
# ============================

import os
import io
import re
import json
from typing import List, Dict, Tuple

import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from groq import Groq


# =====================
# Groq client
# =====================

@st.cache_resource(show_spinner=False)
def get_groq_client():
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("Set GROQ_API_KEY environment variable.")
    return Groq(api_key=api_key)


CHAT_MODEL = "llama-3.3-70b-versatile"  # Groq LLM model


def call_groq_chat(
    messages: List[Dict],
    model: str = CHAT_MODEL,
    temperature: float = 0.0,   # default 0 for deterministic behaviour
) -> str:
    client = get_groq_client()
    chat_completion = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
    )
    return chat_completion.choices[0].message.content


# =====================
# File text extraction
# =====================

def extract_text_from_pdf(file: io.BytesIO) -> str:
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def extract_text_from_docx(file: io.BytesIO) -> str:
    document = Document(file)
    return "\n".join([p.text for p in document.paragraphs])


def extract_text_from_any(uploaded_file) -> str:
    filename = uploaded_file.name.lower()
    content = uploaded_file.read()
    bio = io.BytesIO(content)

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(bio)
    if filename.endswith(".docx"):
        return extract_text_from_docx(bio)
    if filename.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")

    # fallback try utf-8
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""


# =====================
# Core AI logic (skills, roadmap)
# =====================

def parse_json_block(raw: str) -> dict:
    """
    Try hard to pull a JSON object from LLM output.
    """
    text = raw.strip()

    # remove ```json ``` fences if present
    text = re.sub(r"```json", "", text, flags=re.IGNORECASE)
    text = text.replace("```", "").strip()

    # direct JSON?
    try:
        return json.loads(text)
    except Exception:
        pass

    # try to find {...} block
    m = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return {}

    return {}


def extract_skills(resume_text: str, jd_text: str) -> Tuple[List[str], List[str]]:
    """
    Use Groq to extract skills from resume and JD into two clean lists.
    Deterministic + stricter rules to avoid guessing/renaming skills.
    """
    system_prompt = """
You are an expert technical recruiter.

GOAL:
From the given RESUME and JOB DESCRIPTION, extract two clean JSON arrays:
- resume_skills: unique technical and soft skills actually present in the resume text.
- jd_skills: unique skills explicitly or implicitly required in the JD.

STRICT RULES:
- DO NOT guess skills that are not clearly mentioned in the text.
- DO NOT merge or split skills arbitrarily.
  - Example: keep "Python" as "Python", not "Python programming".
  - Keep "Natural Language Processing (NLP)" as written.
- DO NOT change capitalization except normalizing names (e.g., 'python' -> 'Python').
- Each array must contain plain strings only (no nested lists/objects).
- Remove duplicates.

Return ONLY valid JSON in this exact shape:
{
  "resume_skills": ["Python", "Pandas", ...],
  "jd_skills": ["GenAI", "RAG", ...]
}
"""
    user_content = f"RESUME:\n{resume_text}\n\nJOB DESCRIPTION:\n{jd_text}"
    raw = call_groq_chat(
        [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        temperature=0,  # deterministic extraction
    )

    data = parse_json_block(raw)
    resume_skills = data.get("resume_skills", []) or data.get("resumeSkills", [])
    jd_skills = data.get("jd_skills", []) or data.get("job_skills", [])

    # make sure they are flat strings
    resume_skills = [str(s).strip() for s in resume_skills if str(s).strip()]
    jd_skills = [str(s).strip() for s in jd_skills if str(s).strip()]

    return resume_skills, jd_skills


def compute_skill_gaps(resume_skills: List[str], jd_skills: List[str]) -> List[str]:
    resume_set = {s.lower() for s in resume_skills}
    gaps = [s for s in jd_skills if s.lower() not in resume_set]
    return gaps


def compute_resume_score(
    resume_skills: List[str],
    jd_skills: List[str],
) -> Tuple[float, int, int]:
    """
    Simple coverage score: how many JD skills are present in resume.
    Score = matched / total JD skills * 100.
    """
    if not jd_skills:
        return 0.0, 0, 0

    resume_set = {s.lower() for s in resume_skills}
    jd_set = {s.lower() for s in jd_skills}
    overlap = resume_set.intersection(jd_set)

    matched = len(overlap)
    total_required = len(jd_set)
    score = (matched / total_required) * 100 if total_required > 0 else 0.0
    return round(score, 1), matched, total_required


def interpret_resume_score(score: float) -> str:
    if score >= 80:
        return "Excellent match – your resume already fits this JD very well."
    if score >= 60:
        return "Good match – just a few targeted improvements needed."
    if score >= 40:
        return "Okay match – you need to strengthen several skills."
    return "Weak match – there are significant gaps to work on."


def estimate_learning_days(
    missing_skills: List[str],
    experience_level: str,
    daily_hours: float,
    target_timeline_days: int,
) -> Dict[str, int]:
    """
    Estimate focused days per skill INSIDE the overall target timeline.
    We spread the total days across all skills and adjust by experience level
    and daily study hours. Skills are learned in parallel.
    """
    if not missing_skills:
        return {}

    n = len(missing_skills)

    # Experience factor: fresher needs a bit more time per skill
    level = experience_level.lower()
    if "fresher" in level or "intern" in level or "0-1" in level:
        level_factor = 1.2
    elif "mid" in level or "2-5" in level:
        level_factor = 1.0
    else:
        level_factor = 0.8

    # Hours factor: more hours/day = fewer calendar days per skill
    if daily_hours <= 0:
        daily_hours = 1.0
    hours_factor = 3.0 / daily_hours  # 3h/day baseline
    hours_factor = max(0.5, min(hours_factor, 2.0))

    base_equal = max(3.0, target_timeline_days / n)

    result: Dict[str, int] = {}
    for skill in missing_skills:
        days = base_equal * level_factor * hours_factor
        days_int = int(round(days))
        days_int = max(3, min(days_int, target_timeline_days))
        result[skill] = days_int

    return result


def build_learning_roadmap(
    skill_days: Dict[str, int],
    target_timeline_days: int,
) -> str:
    if not skill_days:
        return "No missing skills detected. You already match the JD requirements very well! 🎉"

    skill_lines = "\n".join(
        [f"- {s}: {d} focused days (approx.)" for s, d in skill_days.items()]
    )
    prompt = f"""
Design a realistic weekly learning roadmap for this candidate.

VERY IMPORTANT RULES:
- Total overall timeline should be about **{target_timeline_days} calendar days**.
- The candidate learns multiple skills in **parallel**, not one-by-one.
- Do NOT assign more than **7 calendar days** to any week.
- Use the skill estimates below as *focus days* inside that global plan,
  not as separate, full-time blocks.

Skills and approximate focused days:
{skill_lines}

Output a markdown plan with:
- Week-by-week sections (Week 1, Week 2, ...).
- At the top, add a short note explaining that skills are learned in parallel,
  so the total duration is ~{target_timeline_days} days, not the sum of all per-skill days.
- For each week:
  - Goals
  - Which skills are in focus
  - A small project or task
  - A short **Resources** list with 3–6 items, mixing:
    - video or online course
    - book or chapter
    - blog / documentation
    - optional research paper or GitHub repo if relevant.
"""
    return call_groq_chat(
        [
            {
                "role": "system",
                "content": "You are a senior technical mentor and career coach.",
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0,   # stable structure, content may still vary by skills
    )


# ============================
# LLM video / course / book recommendations
# ============================

def recommend_videos_llm(skill: str, video_type: str) -> dict:
    """
    LLM-based dynamic YouTube + Courses + Books + GitHub recommendations.

    video_type options:
    - "short"  → short videos < 10 minutes
    - "long"   → deep-dive videos (1–3 hours)
    - "course" → full playlist / course
    - "books"  → books, PDFs, GitHub, research papers

    IMPORTANT:
    - We do NOT guess exact YouTube video IDs.
    - We instead generate YouTube SEARCH URLs like:
      https://www.youtube.com/results?search_query=freecodecamp+python+for+beginners
      so links are always valid and the user gets relevant results.
    """
    prompt = f"""
You are an AI mentor trained to recommend the BEST learning resources for tech skills.

Skill: "{skill}"
Requested resource type: "{video_type}"

### GLOBAL RULES

- All URLs you output MUST be valid and clickable.
- For YouTube, DO NOT guess exact video IDs.
  - INSTEAD, ALWAYS use **YouTube search URLs** in this format:
    - "https://www.youtube.com/results?search_query=<keywords-joined-with-+>"
  - Example:
    - "https://www.youtube.com/results?search_query=freecodecamp+python+for+beginners"
    - "https://www.youtube.com/results?search_query=krish+naik+machine+learning+full+course"
- Prefer high-quality, popular channels and platforms, such as (depending on the skill):
  - **Programming / AI / Data / GenAI / ML / DL / RAG:**
    - freeCodeCamp, Krish Naik, Codebasics, Sentdex, DeepLearningAI, HuggingFace,
      StatQuest, Corey Schafer, Nicholas Renotte, AssemblyAI, CodeWithHarry,
      Andrew Ng, Deeplearning.ai, Microsoft Developer, Google Cloud Tech.
  - **System Design / Backend / DevOps / Cloud:**
    - Gaurav Sen, TechWorld with Nana, KodeKloud, Hussein Nasser, AWS, Azure,
      Google Cloud, IBM Technology, System Design Interview.
  - **Blockchain / Web3:**
    - Dapp University, Patrick Collins, EatTheBlocks, Moralis Web3, Chainlink,
      Smart Contract Programmer, freeCodeCamp blockchain playlists.
  - **Frontend / Fullstack / JavaScript / React:**
    - Web Dev Simplified, Traversy Media, Fireship, The Net Ninja,
      Codevolution, freeCodeCamp.
  - **Databases / Data Engineering / Big Data:**
    - Data Engineer One, Luke Barousse, Seattle Data Guy, freeCodeCamp, Codebasics.

- Make resources strongly relevant to the skill: "{skill}".
- Avoid very old content if possible (prefer 2020+).

### BEHAVIOUR

1. If video_type == "short":
   - Return 5 short YouTube video search URLs.
   - Each item should be:
     - title: clear topic name, like "10-minute crash intro to Transformers – freeCodeCamp (YouTube search)"
     - url: a YouTube search URL with keywords for short / quick intro.
     - type: "short"

2. If video_type == "long":
   - Return 3 long-form YouTube **search URLs** for workshops / crash courses (1–3 hours).
   - type: "long"

3. If video_type == "course":
   - Return 3 full playlists or course searches (YouTube or platforms).
   - For non-YouTube (Coursera, Udemy, etc.), use their normal URL if you know it
     OR a Google search URL like:
     "https://www.google.com/search?q=<course-name-+skill>+course"
   - type: "course"

4. If video_type == "books":
   - Return:
     - 3 books (could be Amazon / O'Reilly / PDF search) → type: "book"
     - 3 GitHub repositories or project repos → type: "github"
     - 2–3 documentation or official guides → type: "article"
     - 1 research paper or survey (if relevant) → type: "paper"
   - Links can be real docs / GitHub / arXiv, or at least valid Google search URLs.

### RETURN FORMAT

Return ONLY valid JSON like this:

{{
  "resources": [
    {{
      "title": "Short intro to Transformers – freeCodeCamp (YouTube search)",
      "url": "https://www.youtube.com/results?search_query=freecodecamp+transformers+intro+10+minutes",
      "type": "short"
    }}
  ]
}}
    """

    raw = call_groq_chat(
        [
            {
                "role": "system",
                "content": "You are a senior mentor providing curated, accurate, high-quality resources.",
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,   # some diversity ok
    )

    data = parse_json_block(raw)
    if not isinstance(data, dict):
        return {}
    return data


# =====================
# Resume & CV generation (Groq)
# =====================

def generate_ats_resume(
    resume_text: str,
    jd_text: str,
    role: str,
    prompt_override: str = "",
) -> str:
    base_instruction = f"""
Create an ATS-friendly, professional RESUME for the following candidate.

ROLE APPLIED FOR: {role}

Use information from the EXISTING RESUME and align it with the JOB DESCRIPTION.

HEADER RULES (MUST FOLLOW EXACTLY):
- First line: FULL NAME in UPPERCASE.
- Second line: "City, Country | Phone | Email | LinkedIn | GitHub" (include only fields available in the source).

SECTION ORDER (MUST KEEP THIS ORDER EVEN IF USER GIVES EXTRA PREFERENCES):
1. **Career Objective**
2. **Skills**
3. **Projects**
4. **Experience** (or **Internships** if fresher)
5. **Education**
6. **Certifications**
7. **Achievements**
8. **Additional Information**

FORMATTING RULES:
- Use clear markdown-style headings with the section titles above.
- Use bullet points with action verbs and quantified impact where possible.
- Optimize keywords to match the JD but do not invent fake experience.
- Keep formatting simple (no tables or images) so ATS can parse.
- 1–2 pages for junior, up to 3 pages for senior.

You MUST keep the header and section order exactly as specified.
"""
    if prompt_override:
        base_instruction += (
            f"\n\nExtra user preferences (respect these BUT keep the section order above):\n"
            f"{prompt_override}\n"
        )

    user_content = f"EXISTING RESUME:\n{resume_text}\n\nJOB DESCRIPTION:\n{jd_text}"
    return call_groq_chat(
        [
            {
                "role": "system",
                "content": "You are an expert resume writer and ATS optimization specialist.",
            },
            {"role": "user", "content": base_instruction + "\n\n" + user_content},
        ],
        temperature=0,
    )


def generate_cv(
    resume_text: str,
    jd_text: str,
    role: str,
    prompt_override: str = "",
) -> str:
    base_instruction = f"""
Create a detailed, professional CV (Curriculum Vitae) for the candidate.

ROLE APPLIED FOR: {role}

Use information from the EXISTING RESUME and JOB DESCRIPTION.

HEADER RULES (MUST FOLLOW EXACTLY):
- First line: FULL NAME in UPPERCASE.
- Second line: "City, Country | Phone | Email | LinkedIn | GitHub" (include only fields available in the source).

SECTION ORDER (MUST KEEP THIS ORDER EVEN IF USER GIVES EXTRA PREFERENCES):
1. **Career Objective**
2. **Skills**
3. **Projects**
4. **Experience** (or **Internships** if fresher)
5. **Education**
6. **Certifications**
7. **Publications** (if any)
8. **Achievements**
9. **Extra-curricular**
10. **Additional Information**

FORMATTING RULES:
- More detailed than the resume (2–4 pages allowed).
- Use the headings above in this exact order.
- Use bullet points with action verbs and quantified impact.
- Plain text / markdown, ATS-friendly (no tables, no images).
"""
    if prompt_override:
        base_instruction += (
            f"\n\nExtra user preferences (respect these BUT keep the section order above):\n"
            f"{prompt_override}\n"
        )

    user_content = f"EXISTING RESUME:\n{resume_text}\n\nJOB DESCRIPTION:\n{jd_text}"
    return call_groq_chat(
        [
            {
                "role": "system",
                "content": "You are an expert CV writer and career coach.",
            },
            {"role": "user", "content": base_instruction + "\n\n" + user_content},
        ],
        temperature=0.5,
    )


def generate_professional_prompt(raw_prompt: str, context_hint: str = "") -> str:
    prompt = f"""
The user wants to modify their RESUME and CV but doesn't know how to write a good prompt.

Their raw request:
\"\"\"{raw_prompt}\"\"\"


Context: {context_hint}

Rewrite this as a single, clear, professional prompt that an LLM can follow to regenerate an ATS-friendly resume AND CV.
You MUST tell the model to KEEP:
- The fixed header (name + contact line)
- The section order used in the existing resume/CV (Career Objective, Skills, Projects, Experience/Internships, Education, Certifications, Achievements, etc.).

Focus on:
- ordering of sections (e.g., projects on top, education last)
- tone and style (concise, impact-focused)
- layout or length preferences.

Return ONLY the improved prompt text.
"""
    return call_groq_chat(
        [
            {
                "role": "system",
                "content": "You are an expert prompt engineer for resume optimization tasks.",
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
    )


# ---------- DOCX helpers ----------

def text_to_docx_bytes(text: str) -> io.BytesIO:
    """
    Convert plain/markdown-ish text into a .docx in memory.
    """
    doc = Document()
    for line in text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")  # blank line
        else:
            doc.add_paragraph(line.strip())
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ============================
# Streamlit UI
# ============================

def main():
    st.set_page_config(
        page_title="Gen-AI Skill Gap & Resume Coach",
        page_icon="🧠",
        layout="wide",
    )

    st.title("🧠 Gen-AI Skill Gap Predictor + Resume & CV Builder (Groq)")

    # Make steps super visible
    st.markdown(
        """
### 🔢 Steps in this app

**1️⃣ Upload & Analyze** – upload resume + JD, extract skills & gaps  
**2️⃣ Learning Plan & AI Resources** – generate roadmap + AI video/course/book suggestions  
**3️⃣ Resume & CV** – ATS resume, detailed CV, and prompt-based customization  
"""
    )

    st.markdown(
        """
This app:
- Finds what skills you already have vs what the **role** needs  
- Gives a **Resume–JD match score** so you know how close you are  
- Estimates how long to learn each missing skill (inside your total timeline)  
- Builds a **personal learning roadmap** with weekly plan + resources  
- Generates an **ATS-friendly resume + CV**, and lets you change structure using your own prompt or an auto-improved prompt.  
- Uses an **LLM-based recommendation engine** to suggest short videos, long videos, full courses, and books/GitHub/docs for each missing skill.
"""
    )

    # Sidebar
    st.sidebar.header("⚙️ Candidate Settings")
    role = st.sidebar.text_input(
        "Role you're applying for",
        placeholder="e.g., GenAI Engineer",
    )

    experience_level = st.sidebar.selectbox(
        "Experience level",
        ["Fresher / 0-1 years", "Mid-level / 2-5 years", "Senior / 6+ years"],
    )

    st.sidebar.markdown("#### Daily study time")
    hours = st.sidebar.number_input(
        "Hours",
        min_value=0,
        max_value=12,
        value=3,
        step=1,
    )
    minutes = st.sidebar.selectbox("Minutes", [0, 15, 30, 45], index=0)
    daily_hours = hours + minutes / 60.0

    target_timeline = st.sidebar.slider(
        "Target learning timeline (days)",
        min_value=1,
        max_value=120,
        value=60,
        step=1,
    )

    st.sidebar.markdown("---")
    st.sidebar.write("Make sure `GROQ_API_KEY` is set in your environment.")

    tab1, tab2, tab3 = st.tabs(
        ["1️⃣ Upload & Analyze", "2️⃣ Learning Plan & AI Resources", "3️⃣ Resume & CV"]
    )

    # Session state
    if "resume_text" not in st.session_state:
        st.session_state.resume_text = ""
    if "jd_text" not in st.session_state:
        st.session_state.jd_text = ""
    if "resume_skills" not in st.session_state:
        st.session_state.resume_skills = []
    if "jd_skills" not in st.session_state:
        st.session_state.jd_skills = []
    if "gaps" not in st.session_state:
        st.session_state.gaps = []
    if "skill_days" not in st.session_state:
        st.session_state.skill_days = {}
    if "roadmap" not in st.session_state:
        st.session_state.roadmap = ""
    if "generated_resume" not in st.session_state:
        st.session_state.generated_resume = ""
    if "generated_cv" not in st.session_state:
        st.session_state.generated_cv = ""
    if "improved_prompt" not in st.session_state:
        st.session_state.improved_prompt = ""
    if "resume_score" not in st.session_state:
        st.session_state.resume_score = 0.0
    if "resume_score_comment" not in st.session_state:
        st.session_state.resume_score_comment = ""

    # Tab 1
    with tab1:
        st.subheader("📄 STEP 1: Upload Resume & JD")

        st.markdown("**A. Upload Resume (any type: PDF, DOCX, TXT, ...)**")
        resume_file = st.file_uploader(
            "Upload your resume file", type=None, key="resume_uploader"
        )

        if resume_file is not None:
            st.session_state.resume_text = extract_text_from_any(resume_file)

        st.text_area(
            "Extracted resume text (editable before analysis)",
            value=st.session_state.resume_text,
            height=220,
        )

        st.markdown("**B. Provide Job Description (JD)**")
        jd_option = st.radio(
            "How do you want to give the JD?",
            ["Paste text", "Upload file"],
            horizontal=True,
        )

        if jd_option == "Paste text":
            jd_text_area = st.text_area(
                "Paste JD here",
                value=st.session_state.jd_text,
                height=220,
            )
            st.session_state.jd_text = jd_text_area
        else:
            jd_file = st.file_uploader(
                "Upload JD file", type=None, key="jd_uploader"
            )
            if jd_file is not None:
                st.session_state.jd_text = extract_text_from_any(jd_file)

            st.text_area(
                "Extracted JD text (editable before analysis)",
                value=st.session_state.jd_text,
                height=220,
            )

        if st.button("🔍 Analyze Skills & Gaps"):
            if not st.session_state.resume_text or not st.session_state.jd_text:
                st.error("Please provide both resume and JD.")
            else:
                with st.spinner("Using Groq to extract skills..."):
                    resume_skills, jd_skills = extract_skills(
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                    )
                st.session_state.resume_skills = resume_skills
                st.session_state.jd_skills = jd_skills
                st.session_state.gaps = compute_skill_gaps(
                    resume_skills,
                    jd_skills,
                )

                # Compute resume–JD score
                score, matched, total_required = compute_resume_score(
                    resume_skills,
                    jd_skills,
                )
                st.session_state.resume_score = score
                st.session_state.resume_score_comment = interpret_resume_score(score)

                st.success("Skill analysis complete! ✅")

        # Show results if available
        if st.session_state.resume_skills or st.session_state.jd_skills:
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Resume skills", len(st.session_state.resume_skills))
            with col2:
                st.metric("JD skills (role needs)", len(st.session_state.jd_skills))
            with col3:
                st.metric("Missing skills (gaps)", len(st.session_state.gaps))

            if st.session_state.jd_skills:
                st.markdown(
                    f"### 📊 Resume–JD fit score: "
                    f"**{st.session_state.resume_score:.1f}%**"
                )
                st.write(st.session_state.resume_score_comment)

            st.markdown("### ✅ Skills in your resume")
            if st.session_state.resume_skills:
                st.markdown(", ".join(st.session_state.resume_skills))
            else:
                st.write("_No skills detected from resume text._")

            st.markdown("### 🎯 Skills required by the JD (role wants)")
            if st.session_state.jd_skills:
                st.markdown(", ".join(st.session_state.jd_skills))
            else:
                st.write("_No skills detected from JD._")

            st.markdown("### ⚠️ Missing skills (gaps)")
            if st.session_state.gaps:
                st.markdown(", ".join(st.session_state.gaps))
            else:
                st.write("_No gaps detected – great match!_")

    # Tab 2
    with tab2:
        st.subheader("📆 STEP 2: Learning Roadmap & AI Resources")

        if not st.session_state.jd_skills and not st.session_state.gaps:
            st.info("First run **Analyze Skills & Gaps** in STEP 1.")
        else:
            if st.button("📅 Generate Learning Roadmap"):
                st.session_state.skill_days = estimate_learning_days(
                    st.session_state.gaps,
                    experience_level,
                    daily_hours,
                    target_timeline,
                )
                with st.spinner("Asking Groq to design your roadmap..."):
                    st.session_state.roadmap = build_learning_roadmap(
                        st.session_state.skill_days,
                        target_timeline,
                    )

            if st.session_state.skill_days:
                st.markdown(
                    f"### ⏱ Approximate focused days per missing skill "
                    f"(inside your {target_timeline}-day plan, skills learned in parallel)"
                )
                for skill, days in st.session_state.skill_days.items():
                    st.write(f"- **{skill}** → ~{days} days")

            if st.session_state.roadmap:
                st.markdown("### 📚 Your learning roadmap")
                st.markdown(st.session_state.roadmap)

            st.markdown("---")
            st.subheader("🎥 AI-Generated Video, Course & Book Recommendations")

            if not st.session_state.gaps:
                st.info("No missing skills ⇒ no recommendations needed.")
            else:
                selected_skill = st.selectbox(
                    "Pick a missing skill to get resources for",
                    st.session_state.gaps,
                )

                video_choice = st.radio(
                    "Select resource type:",
                    [
                        "Short Videos (< 10 min)",
                        "Long Videos (1–3 hours)",
                        "Full Courses / Playlists",
                        "Books / GitHub / Docs",
                    ],
                    horizontal=True,
                )

                if st.button("🔎 Get AI Recommendations"):
                    with st.spinner("Asking Groq LLM for best resources..."):
                        if video_choice.startswith("Short"):
                            data = recommend_videos_llm(selected_skill, "short")
                        elif video_choice.startswith("Long"):
                            data = recommend_videos_llm(selected_skill, "long")
                        elif video_choice.startswith("Full Courses"):
                            data = recommend_videos_llm(selected_skill, "course")
                        else:
                            data = recommend_videos_llm(selected_skill, "books")

                    st.markdown("### 📚 AI-Recommended Resources")
                    if not data or "resources" not in data:
                        st.error(
                            "No resources found or parsing failed. You can try again."
                        )
                    else:
                        for item in data["resources"]:
                            title = item.get("title", "Untitled")
                            url = item.get("url", "#")
                            rtype = item.get("type", "resource")
                            st.markdown(
                                f"- **{title}**  \n"
                                f"[{url}]({url})  \n"
                                f"_Type_: `{rtype}`"
                            )

    # Tab 3
    with tab3:
        st.subheader("📝 STEP 3: Resume & CV Generator")

        st.markdown(
            """
**Flow:**

1. Generate first version of ATS **Resume + CV** from your current resume + JD  
2. If you don't like the structure, describe changes in your own words  
3. Use:
   - 🟢 _Regenerate using my prompt (direct)_  
   - 🧠 _Auto-improve my prompt, then regenerate_ (for people with communication issues / not sure how to write prompts)
"""
        )

        if st.button("✨ Generate initial ATS Resume + CV"):
            if not st.session_state.resume_text or not st.session_state.jd_text:
                st.error("Please provide both resume and JD in STEP 1.")
            else:
                with st.spinner("Generating ATS-friendly resume and CV with Groq..."):
                    st.session_state.generated_resume = generate_ats_resume(
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                        role or "Software Engineer",
                    )
                    st.session_state.generated_cv = generate_cv(
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                        role or "Software Engineer",
                    )

        if st.session_state.generated_resume or st.session_state.generated_cv:
            col_r, col_c = st.columns(2)

            with col_r:
                st.markdown("### 📄 Resume (ATS-friendly)")
                st.text_area(
                    "Resume text",
                    value=st.session_state.generated_resume,
                    height=320,
                )
                if st.session_state.generated_resume:
                    resume_docx = text_to_docx_bytes(
                        st.session_state.generated_resume
                    )
                    st.download_button(
                        "⬇️ Download Resume (.docx)",
                        data=resume_docx,
                        file_name="resume_ats.docx",
                        mime=(
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    )

            with col_c:
                st.markdown("### 📄 CV (detailed)")
                st.text_area(
                    "CV text",
                    value=st.session_state.generated_cv,
                    height=320,
                )
                if st.session_state.generated_cv:
                    cv_docx = text_to_docx_bytes(st.session_state.generated_cv)
                    st.download_button(
                        "⬇️ Download CV (.docx)",
                        data=cv_docx,
                        file_name="cv_detailed.docx",
                        mime=(
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    )

        st.markdown("---")
        st.subheader("🎯 Customize Resume & CV structure via prompt")

        user_change_prompt = st.text_area(
            "Type how you want to change the RESUME + CV structure / order / style",
            placeholder=(
                "e.g., Put Projects section at the top, keep Education at last, "
                "short summary, highlight GenAI + RAG and internships..."
            ),
            height=150,
        )

        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("♻️ Regenerate using my prompt (direct)"):
                if not user_change_prompt.strip():
                    st.error("Please describe at least one change.")
                else:
                    with st.spinner("Regenerating resume & CV using your prompt..."):
                        st.session_state.generated_resume = generate_ats_resume(
                            st.session_state.resume_text,
                            st.session_state.jd_text,
                            role or "Software Engineer",
                            prompt_override=user_change_prompt,
                        )
                        st.session_state.generated_cv = generate_cv(
                            st.session_state.resume_text,
                            st.session_state.jd_text,
                            role or "Software Engineer",
                            prompt_override=user_change_prompt,
                        )

        with col_b:
            if st.button("🧠 Auto-improve my prompt, then regenerate"):
                if not user_change_prompt.strip():
                    st.error("Please describe at least one change.")
                else:
                    with st.spinner(
                        "Improving your prompt and regenerating resume & CV..."
                    ):
                        improved = generate_professional_prompt(
                            user_change_prompt,
                            context_hint=f"Role: {role}, Experience: {experience_level}",
                        )
                        st.session_state.improved_prompt = improved
                        st.session_state.generated_resume = generate_ats_resume(
                            st.session_state.resume_text,
                            st.session_state.jd_text,
                            role or "Software Engineer",
                            prompt_override=improved,
                        )
                        st.session_state.generated_cv = generate_cv(
                            st.session_state.resume_text,
                            st.session_state.jd_text,
                            role or "Software Engineer",
                            prompt_override=improved,
                        )

        if st.session_state.improved_prompt:
            st.markdown(
                "### ✍️ Auto-generated professional prompt (used for last regeneration)"
            )
            st.code(st.session_state.improved_prompt, language="markdown")


if __name__ == "__main__":
    main()
